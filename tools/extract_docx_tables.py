#!/usr/bin/env python3
"""
Extract tables from a DOCX file.

Usage:
  python3 tools/extract_docx_tables.py path/to/source.docx
  python3 tools/extract_docx_tables.py path/to/source.docx --tables 0,2-4 --out tmp/tables.md
  python3 tools/extract_docx_tables.py path/to/source.docx --tables 5 --format csv --out tmp/table-5.csv

Table indexes are zero-based, matching python-docx's table order.
"""

from __future__ import annotations

import argparse
import csv
import sys
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
except ImportError:  # pragma: no cover - depends on local environment
    Document = None  # type: ignore[assignment]


def parse_table_selection(raw: str, table_count: int) -> list[int]:
    if raw.lower() == "all":
        return list(range(table_count))

    selected: set[int] = set()

    for part in raw.split(","):
        part = part.strip()
        if not part:
            continue

        if "-" in part:
            start_raw, end_raw = part.split("-", 1)
            start = int(start_raw)
            end = int(end_raw)
            if end < start:
                raise ValueError(f"invalid table range: {part}")
            selected.update(range(start, end + 1))
        else:
            selected.add(int(part))

    indexes = sorted(selected)
    for index in indexes:
        if index < 0 or index >= table_count:
            raise ValueError(f"table index out of range: {index}; file has {table_count} table(s)")

    return indexes


def table_to_matrix(table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip().replace("\r", "\n") for cell in row.cells])
    return rows


def markdown_escape_cell(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", "<br>")


def render_markdown_table(matrix: list[list[str]]) -> str:
    if not matrix:
        return "_Empty table_"

    width = max((len(row) for row in matrix), default=0)
    normalized = [row + [""] * (width - len(row)) for row in matrix]

    header = [markdown_escape_cell(cell) for cell in normalized[0]]
    separator = ["---"] * width
    body = [[markdown_escape_cell(cell) for cell in row] for row in normalized[1:]]

    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]

    for row in body:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def render_markdown(docx_path: Path, table_indexes: Iterable[int], tables) -> str:
    parts = [
        f"# Extracted tables",
        "",
        f"- Source: `{docx_path.as_posix()}`",
        f"- Table indexes: zero-based",
        "",
    ]

    for index in table_indexes:
        matrix = table_to_matrix(tables[index])
        parts.append(f"## Table {index}")
        parts.append("")
        parts.append(render_markdown_table(matrix))
        parts.append("")

    return "\n".join(parts).rstrip() + "\n"


def write_csv(path: Path, matrix: list[list[str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerows(matrix)


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract DOCX tables to Markdown or CSV.")
    parser.add_argument("docx_path", help="Path to the DOCX source file.")
    parser.add_argument(
        "--tables",
        default="all",
        help="Zero-based table indexes to extract, e.g. all, 0, 2-4, or 0,2-4.",
    )
    parser.add_argument(
        "--format",
        choices=("markdown", "csv"),
        default="markdown",
        help="Output format. CSV requires selecting exactly one table.",
    )
    parser.add_argument("--out", help="Output file. Defaults to stdout for Markdown.")
    args = parser.parse_args()

    if Document is None:
        print(
            "extract_docx_tables: missing dependency python-docx. Install with: pip install python-docx",
            file=sys.stderr,
        )
        return 2

    docx_path = Path(args.docx_path)
    if not docx_path.is_file():
        print(f"extract_docx_tables: file not found: {docx_path}", file=sys.stderr)
        return 1

    document = Document(str(docx_path))
    tables = document.tables

    try:
        table_indexes = parse_table_selection(args.tables, len(tables))
    except ValueError as exc:
        print(f"extract_docx_tables: {exc}", file=sys.stderr)
        return 1

    if args.format == "csv":
        if len(table_indexes) != 1:
            print("extract_docx_tables: CSV output requires exactly one selected table", file=sys.stderr)
            return 1
        if not args.out:
            print("extract_docx_tables: CSV output requires --out", file=sys.stderr)
            return 1

        matrix = table_to_matrix(tables[table_indexes[0]])
        write_csv(Path(args.out), matrix)
        print(f"extract_docx_tables: wrote {args.out}", file=sys.stderr)
        return 0

    output = render_markdown(docx_path, table_indexes, tables)

    if args.out:
        out_path = Path(args.out)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(output, encoding="utf-8")
        print(f"extract_docx_tables: wrote {args.out}", file=sys.stderr)
    else:
        print(output, end="")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
